# 导入库
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

# 新建空白文档
doc1 = Document()

# 新增文档标题
doc1.add_heading('如何使用 Python 创建和操作 Word',0)

# 创建段落描述
doc1.add_paragraph(' Word 文档在我们现在的生活和工作中都用的比较多，我们平时都使用 wps 或者 office 来对 Word 进行处理，可能没想过它可以用 Python 生成，下面我们就介绍具体如何操作……')

# 创建一级标题
doc1.add_heading('安装 python-docx 库',1)

# 创建段落描述
doc1.add_paragraph('现在开始我们来介绍如何安装 python-docx 库，具体需要以下两步操作：')

# 创建二级标题
doc1.add_heading('第一步：安装 Python',2)

# 创建段落描述
doc1.add_paragraph('在python官网下载python安装包进行安装。')

# 创建三级标题
doc1.add_heading('第二步：安装 python-docx 库',3)

# 创建段落描述
doc1.add_paragraph('window下win+R输入CMD打开命令行，输入pip install python-docx即可下载。')

# 保存文件
doc1.save('word2.docx')
